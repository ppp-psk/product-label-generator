#Version 1.2 Date Feb 28, 2021
# - Read data from .xlsx and output .docx

import sys #For exit program
import datetime #Create output file name
import time #Program sleep function
from docx import Document #Create .docx file
from docx.shared import Length #doc layout
from docx.shared import Pt #doc layout
from docx.shared import Mm #doc layout
from docx.enum.text import WD_ALIGN_PARAGRAPH #doc layout
import pandas as pd #Datatable processing

#CHANGE INPUT FILE HERE
inputFilename = "product_input.xlsx"

global cellPointer  #There are four cell per page (use in writeProductDataToTable)
global table #Temporary variable for current avaliable table (use in writeProductDataToTable)
table = None
cellPointer = 0

#Convert quantity count into โหล and ชิ้น
def generateQuantityString(quantity):
    dozen = int(quantity/12)
    remaining = int(quantity%12)
    resultString = ""
    if dozen > 0:
        resultString = str(dozen) + " โหล"
    
    if remaining > 0:
        resultString = resultString + " " + str(remaining) + " ชิ้น"
    
    return resultString

#Insert product label into cells in tables
def writeProductDataToTable(dozenCount, modelName, modelColor, modelSize , modelQuantity):
    global cellPointer
    global table
    number = str(dozenCount)
    model = modelName
    color = modelColor
    size = str(modelSize)
    quantityString = generateQuantityString(modelQuantity)
    
	#Add new table into blank page
    if(cellPointer % 4 == 0):
		#Move to new page, when table in current page is full
        if(table != None):
            document.add_page_break()
        table = document.add_table(rows=2, cols=2)
        cellPointer = 0
    
	#a List of component for one time setting 
    settings = list()
    rowIndex = int(cellPointer/2) 
    columnIndex = cellPointer % 2
    
    cells  = table.rows[rowIndex].cells

    p = cells[columnIndex].add_paragraph('(ลำดับที่)')
    p.add_run(' ' + number).bold = True
    settings.append(p)

    if(len(model) < 10):
        p = cells[columnIndex].add_paragraph('(รุ่น)')
        p.add_run(' ' + model).bold = True
        settings.append(p)
    else:
        p = cells[columnIndex].add_paragraph()
        p.add_run(' ' + model).bold = True
        settings.append(p)

    p = cells[columnIndex].add_paragraph('(สี)')
    p.add_run(' ' + color).bold = True
    settings.append(p)

    p = cells[columnIndex].add_paragraph('(ขนาด)')
    p.add_run(' ' + size).bold = True
    settings.append(p)

    p = cells[columnIndex].add_paragraph('(จำนวน)')
    p.add_run(' ' + quantityString).bold = True
    settings.append(p)       

    for text in settings:
        text.alignment  = WD_ALIGN_PARAGRAPH.CENTER
    
    cellPointer = cellPointer + 1
    
	


try:
	product_data = pd.read_excel(inputFilename)

	try:
		document = Document()

		section = document.sections[0]

		section.page_height = Mm(210)
		section.page_width = Mm(120)
		section.left_margin = Mm(13)
		section.right_margin = Mm(10)
		section.top_margin = Mm(10)
		section.bottom_margin = Mm(10)

		style = document.styles['Normal']
		font = style.font
		font.name = 'TH Sarabun New'
		font.size = Pt(20)

		columnList = product_data.columns.tolist()
		startColumnSize = 2
		productRowCount = product_data.shape[0]
		productSizeCount = product_data.shape[1] - startColumnSize

		#Each size in each model
		for i in range(productRowCount):
			productModelName = product_data['model'][i]
			productModelColor = product_data['color'][i]
			for j in range(productSizeCount):
				productModelSize = product_data.columns[(startColumnSize ) + j]
				currentProductSizeCount = product_data[productModelSize][i]
				dozenCount = 1
				while (currentProductSizeCount > 0):
					if(currentProductSizeCount > 12 and currentProductSizeCount < 18):
						writeProductDataToTable(dozenCount, productModelName, productModelColor, productModelSize , currentProductSizeCount)
						dozenCount = dozenCount + 1
						currentProductSizeCount = 0
					elif(currentProductSizeCount >= 12):
						writeProductDataToTable(dozenCount, productModelName, productModelColor, productModelSize , 12)
						currentProductSizeCount = currentProductSizeCount - 12
						dozenCount = dozenCount + 1
					elif(currentProductSizeCount < 12):
						writeProductDataToTable(dozenCount, productModelName, productModelColor, productModelSize , currentProductSizeCount)
						currentProductSizeCount = 0
						dozenCount = dozenCount + 1

		timeNow = datetime.datetime.now()
		document.save( str(timeNow.year)+ str(timeNow.month) + str(timeNow.day)+ '_' + str(timeNow.hour) + "-" + str(timeNow.minute)+ "-" + str(timeNow.second)+ '.docx')
	except  Exception as e:
		#Error for internal code
		print("[Error] Please check following error")
		print(e)
		time.sleep(10)
except:
	#Error for O/I file exception
	print("[Error] Cannot open " + inputFilename)
	time.sleep(10)








